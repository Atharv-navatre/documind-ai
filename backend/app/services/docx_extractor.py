"""
DOCX Text Extractor
===================
Extracts text from Microsoft Word documents using python-docx.
"""

from docx import Document


def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts text from all paragraphs of a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
    
    Returns:
        str: Combined text from all paragraphs
    
    Raises:
        ValueError: If DOCX cannot be opened or read
    """
    try:
        # Open the document
        doc = Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        # Combine all text
        full_text = "\n\n".join(paragraphs)
        
        # Clean up
        full_text = clean_extracted_text(full_text)
        
        if not full_text.strip():
            raise ValueError("DOCX appears to be empty (no extractable text)")
        
        return full_text
        
    except Exception as e:
        if "not a valid" in str(e).lower() or "zip" in str(e).lower():
            raise ValueError(f"Invalid or corrupted DOCX file: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {e}")


def clean_extracted_text(text: str) -> str:
    """
    Cleans extracted text by removing excessive whitespace.
    
    Args:
        text: Raw extracted text
    
    Returns:
        str: Cleaned text
    """
    import re
    
    # Replace multiple newlines with double newline
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Replace multiple spaces with single space
    text = re.sub(r' {2,}', ' ', text)
    
    # Strip leading/trailing whitespace
    text = text.strip()
    
    return text
